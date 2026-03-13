"""
Word文档导出模块
"""
import io
from datetime import datetime
from typing import List, Dict, Any
import json

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE


def export_to_word(
    project_title: str,
    outline_json: str,
    checkpoints: List[Dict[str, Any]],
    format_config: Dict[str, Any] = None
) -> bytes:
    """
    导出标书为Word文档

    Args:
        project_title: 项目标题
        outline_json: 大纲JSON
        checkpoints: 章节内容列表
        format_config: 格式配置 {"target_pages": int, "words_per_page": int}

    Returns:
        Word文档的字节内容
    """
    doc = Document()

    # 默认格式配置
    if format_config is None:
        format_config = {}
    target_pages = format_config.get("target_pages", 50)
    words_per_page = format_config.get("words_per_page", 700)

    # 设置文档标题
    title = doc.add_heading(project_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 添加文档信息（格式设置信息）
    info_paragraph = doc.add_paragraph()
    info_paragraph.add_run(f"目标页数：{target_pages}页 | 每页约{words_per_page}字").font.size = Pt(10)
    info_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_paragraph.space_after = Pt(20)

    # 计算总字数
    total_words = sum(cp.get("word_count", 0) for cp in checkpoints)

    # 添加统计信息
    stats_paragraph = doc.add_paragraph()
    stats_paragraph.add_run(f"本文档共 {len(checkpoints)} 章，约 {total_words} 字").font.size = Pt(10)
    stats_paragraph.space_after = Pt(20)

    # 解析大纲获取章节顺序
    outline = json.loads(outline_json) if outline_json else []
    chapter_order = []

    def extract_chapters(nodes, parent_title=""):
        for node in nodes:
            if node.get("level") == 1:
                parent_title = node.get("title", "")
            if node.get("level") in [2, 3]:
                chapter_order.append({
                    "title": node.get("title"),
                    "level": node.get("level"),
                    "parent": parent_title
                })
            if node.get("children"):
                extract_chapters(node["children"], parent_title)

    extract_chapters(outline)

    # 创建章节内容映射
    content_map = {cp["chapter_title"]: cp["content"] for cp in checkpoints}

    # 添加正文内容
    current_level1 = None
    for chapter in chapter_order:
        title_text = chapter["title"]

        if chapter["level"] == 1:
            # 一级标题
            heading = doc.add_heading(title_text, level=1)
            current_level1 = title_text
        elif chapter["level"] == 2:
            # 二级标题
            if current_level1:
                # 添加父章节标题作为前缀
                pass
            heading = doc.add_heading(title_text, level=2)
        else:
            # 三级标题
            heading = doc.add_heading(title_text, level=3)

        # 添加正文内容
        content = content_map.get(title_text, "")
        if content:
            # 清理内容（移除可能的思考标记）
            content = content.replace("<think>", "").replace("</think>", "")

            # 分割成段落
            paragraphs = content.split("\n\n")
            for para in paragraphs:
                if para.strip():
                    p = doc.add_paragraph(para.strip())
                    # 设置正文格式
                    for run in p.runs:
                        run.font.size = Pt(12)
                        run.font.name = "宋体"

    # 添加页脚
    section = doc.sections[-0] if doc.sections else None
    if section:
        footer = section.footer
        paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        paragraph.text = f"BidAI智能投标系统生成 · {datetime.now().strftime('%Y-%m-%d')}"
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 保存到字节流
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer.getvalue()


def get_default_format() -> Dict[str, Any]:
    """获取默认格式配置"""
    return {
        "h1": {
            "font": "黑体",
            "size": 16,
            "align": "center",
            "bold": True
        },
        "h2": {
            "font": "黑体",
            "size": 14,
            "align": "left",
            "bold": True
        },
        "h3": {
            "font": "黑体",
            "size": 12,
            "align": "left",
            "bold": True
        },
        "body": {
            "font": "宋体",
            "size": 12,
            "lineHeight": 1.5
        },
        "page": {
            "size": "A4",
            "margin": 2.5
        }
    }
